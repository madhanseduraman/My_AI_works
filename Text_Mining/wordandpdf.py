### THIS IS A TEXT EXTRACTION SOLUTION FROM PDF AND WORD FILES. COMPLETELY ABOUT PROCESSING UNSTRUCTURED DATA. EXTRACTS PARAGRAPHS REQUIRED BASED ON THE HEADING WORDS INPUT PROVIDED ###



import pandas as pd
import os
import re
import nltk
import io
import docx
from unidecode import unidecode
from collections import Counter
import sys
from datetime import date



def listtextfiles(samplepath):
    textFiles = []
    #newpdfFiles = []
    for root, dirs, files in os.walk(samplepath):
        for file in files:
            #print(file)
            if file.endswith(".txt"):
                textFiles.append(os.path.join(root, file))
            elif file.endswith(".TXT"):
                textFiles.append(os.path.join(root, file))
    #print(textFiles)
    return(textFiles)


# In[5]:


def readfromText(filename):
    f = open(filename,"r",encoding='utf-8')
    text = f.read()
    f.close()
    return text


# #### Get Table of Contents

# In[6]:

def getcontentspdf(filename):
    toc = []
    reg = []
    f = open(filename,"r",encoding='utf-8')
    for i,line in enumerate(f):
        #print(line)
        if i<=51:
            if "......." in line :
                toc.append(line)
            
            #abc=line.split(' ')
            #if len(abc)==1:
                #toc.append(line)
            else:
                
                se = re.search(r'(^.*(?=\d))',line)
                #se = re.search(r'(\d+$)',line)
                if se:
                    toc.append(se.group(0))
    c4 = []
    for c in toc:
        c1 = re.sub(r'\d+','',c)
        c2 = re.sub(r'[._]','',c1)
        if len(c2)>3:
            c3 = c2.strip()
            c4.append(c3)
    if not c4:
        return('TOC not found')
        
    f.close()
    c4 = [i for i in c4 if i!= '']
    #print(c4)
    return c4


# In[7]:


def getmatchindex(cwordlist):
    kwords = list(keywords.words)
   # print(cwordlist)
    match_ind = []
    match_words = []
    #cwordlist = [item.lower() for item in cwordlist]
    kwords1 = [m.split(' ') for m in kwords]
    l = [j for i in kwords1 for j in i]
    l = list(filter(lambda x:x, map(lambda x:re.sub(r'[^A-Za-z]', '', x), l)))
  
    
    #stops = set(stopwords.words("english"))
    words = [w for w in l if not w.lower() in stops]      
   # print(words)
    for i,cword in enumerate(cwordlist):
        for k in kwords:
            #print(k)
            if k in cword:
               # print("key",k)
                #print("cword",cword)
                
               # if len(cword.split()) <5:
                #if k not in match_words:
                    #print(cword)
                index = i
                match_ind.append(index)
                match_words.append(k)
                
            
    #print("matched index",match_words)
    return match_ind


# In[8]:


def getmatchwords(cwordlist):
    #print(cwordlist)
    #cwordlist=["part a- introduction","program"]
    kwords = list(keywords.words)
    match_ind = []
    match_words = []
       # print(kwords)
    #cwordlist = [item.lower() for item in cwordlist]
    kwords1 = [m.split(' ') for m in kwords]
    l = [j for i in kwords1 for j in i]
    l = list(filter(lambda x:x, map(lambda x:re.sub(r'[^A-Za-z]', '', x), l)))

    #stops = set(stopwords.words("english"))
    words = [w for w in l if not w.lower() in stops]

    
    for i,cword in enumerate(cwordlist):
        for k in kwords:
            #print(k)
            if k in cword:
                #print("key",k)
                #print("cword",cword)
                #if len(cword.split()) <5:
                #if k not in match_words:
                match_words.append(k)
                
    
   # print("matched data")               
    print(match_words)
    return match_words


# #### Identify & Remove Footer information 

# In[9]:


def getfooter(f):
    my = open(f,"r",encoding='utf-8')
    for line in my:
        line = re.sub(r'\d','',line)
        with open(samplepath+'\\dummy.txt','a',encoding='utf-8') as t:
            t.write(line)
    dummyfile = samplepath+'\\dummy.txt'
    newfile =  open(dummyfile,'r',encoding='utf-8')
    lines = newfile.readlines()
    c=Counter(c[0:25] for c in lines if c.strip()) 
    most_com = c.most_common(3)
    com_text = [x[0] for x in most_com]
    #print(com_text)
    my.close()
    newfile.close()
    
    os.remove(dummyfile)
    
    return(com_text)


# In[10]:


def txtoutput (f,content_1,content_2):
    com_text = getfooter(f)
    #print(com_text[2])
   # print("content",content_1)
    #print("content",content_1.lower())
    #print("content",content_1.title())
   # print("llll",content_2)
    #skip = True
    my_file = open(f,"r",encoding='utf-8')
    k=my_file.readlines()
    op_filename = samplepath+'\output.txt'
    op_file = open(op_filename,"a",encoding='utf-8')
#   comment the below code for getting the o/p in a text file       
    op_file.truncate(0)
    flag = 0
    c=0
    block = ''
   
    abc=[]
    #print(k)
    for line in range(0,len(k)):
        
        if content_1.upper() in k[line] and "_______" not in k[line] and \
        not re.search(r'\d+$',k[line]) or content_1.lower() in k[line] and \
        "_______" not in k[line] and not re.search(r'\d+$',k[line]) or \
        content_1.title() in k[line] and "_______" not in k[line] and \
        not re.search(r'\d+$',k[line]) or content_1 in k[line] and "_______" not in k[line] and \
            not re.search(r'\d+$',k[line]):
            #print(content_1)
            if len(k[line].split()) <= 8:
            
                for i in range(line+1,len(k)):
                    
                    if content_2.lower()  in k[i].lower():
                        #print("linecon",content_2.lower())
                        #print("conteline",len(k[i]))
                        #print("first",content_2.lower())
                        
                        
                        if len(k[i]) <= len(content_2)+5:
                           
                            for j in range(line,i):
                                if com_text[0] not in k[j] and com_text[1] not in k[j]:
                                    abc.append(k[j])
                             
                            break

    lmn=[]
   # print(len(abc))
    #print("hai")
    #print(abc)
    if len(abc) <= 1:
        str1=''
    elif len(abc) >500:
        str1 = 'cannot find an extractable content'
    else:
        #print(abc)
        for each1 in abc:
            if len(each1)>3:
                if each1[0].isalnum() and each1[1] not in symbol and each1[2] not in symbol and each1[3] not in symbol:
                    a=each1.rstrip()
                    lmn.append(a)
                        
                elif len(each1.split()) >6:
                   
                    lmn.append("\n"+each1)
                else:
                    lmn.append("\n"+each1)
            else:
               lmn.append("\n"+each1)
                
                
                
        #print(lmn)
        

            
        str1 = ''.join(lmn)
        #str1 = ''.join(abc)
        #wrapper = textwrap.TextWrapper(width=50) 
  
        #word_list = wrapper.wrap(text=str1) 
          
        # Print each line.
       # linebreakage=[]
       # for element in word_list: 
          #  linebreakage.append(element)
            
       # str1=''.join(linebreakage)
       # if c>2:
          #  break
            #break
        #print("--------------------------2nd occurrence",content_1)
    #print(abc)
    #abcde=str1
    #a=abcde.rstrip('')
    #print(a)
    my_file.close()
    op_file.close()
    text = readfromText(op_filename)
    #text = textwrap.fill(text,fix_sentence_endings=True,replace_whitespace=False)
    os.remove(op_filename)
    #print(unidecode(block))
    return unidecode(str1)


# def blocktxtoutput (f,content_1,content_2):
#     com_text = getfooter(f)
#     #skip = True
#     my_file = open(f,"r",encoding='utf-8')
#     op_filename = samplepath+'\output.txt'
#     op_file = open(op_filename,"a",encoding='utf-8')
# #   comment the below code for getting the o/p in a text file       
#     op_file.truncate(0)
#     flag = 0
#     c=0
#     block = ''
#     for line in my_file:
#         if content_1 in line and "....." not in line:
#             c+=1
#             continue
#         if content_2 in line:
#             break
#         block+=line
#         if c>1:
#             break
#         #break
#     #print(block)
#         #print("--------------------------2nd occurrence",content_1)
#     my_file.close()
#     op_file.close()
#     text = block
#     #print(unidecode(text))
#     #text = textwrap.fill(text,fix_sentence_endings=True,replace_whitespace=False)
#     os.remove(op_filename)
#     return unidecode(block)

# #### Extract Targetted Text

# In[86]:


def getpdfoutput(f):
    #print(f)
    splitpd = os.path.splitext(f)[0]
    pdffilename = os.path.basename(splitpd)
    output_list = []
    list_cleaned = []
    list_cleaned1 = []
    content_words = getcontentspdf(f)
    #print("contentword")
    #print(content_words)
    leng  = len(content_words)
    matchin = getmatchindex(content_words)
    matchwor = getmatchwords(content_words)
   # print(matchin)
    print(content_words)
    length=len(content_words)
    for m in matchin:
        a=m+1
        #print(a)
        if a>=length:
            output_list.append(" ")
            #print("is eof")
        else:
            if (leng-m)>1:
                text1 = content_words[m]
                text2 = content_words[m+1]
                #print(text1)
                #print(text2)
                
                text6 = content_words[m].title()
                text7 = content_words[m+1].title()
                text8 = content_words[m].upper()
                text9 = content_words[m+1].upper()
                output = txtoutput(f,text1,text2)
                #print(output)
                if output:
                    #print("1")
                    #print(text1)
                    #print(text2)
                    output_list.append(output)
                else:
                   # print("2")
                    
                    #print(text6)
                    #print(text7)
                    output = txtoutput(f,text6,text7)
                    if output:
                        #print("3")
                        
                        output_list.append(output)
                    else:
                       # print("4")
                        #print(text8)
                        #print(text9)
                        #output_list = [wrapper.wrap(i) for i in output.split('\n') if i != '']
                        output = txtoutput(f,text8,text9)
                        if output:
                           # print("5")
                            output_list.append(output)
                        else:
                            #print("same code with m+2")
                            a=m+2
                           # print("length of match",m)
                           # print("length of content word",length)
                            if a>=length:
                                output_list.append("end of table of content")
                                #print("eof")
                            else:
                               # print("6")
                                text1 = content_words[m]
                                text2 = content_words[m+2]
                                #print(text1)
                                #print(text2)
                                text6 = content_words[m].title()
                                text7 = content_words[m+2].title()
                                text8 = content_words[m].upper()
                                text9 = content_words[m+2].upper()
                                output = txtoutput(f,text1,text2)
                                if output:
                                   # print("7")
                                    #print(text1)
                                    #print(text2)
                                    output_list.append(output)
                                else:
                                   # print("8")
                                    output = txtoutput(f,text6,text7)
                                    if output:
                                       # print("9")
                                        #print(text6)
                                        #print(text7)
                                        output_list.append(output)
                                    else:
                                       # print("10")
                                        #print(text8)
                                        #print(text9)
                                        output = txtoutput(f,text8,text9)
                                        output_list.append(output)
                        
   # print("outputlist--------",output_list)
   # out = str(output_list).split()
   # print("output________",len(out))
    #if len(out)<2000:
    print(output_list)
    return output_list
   # else:
	#    return ''


	

# In[88]:
def extractFile(pathFromCmd):
    global samplepath
    samplepath=pathFromCmd
    #print(samplepath)
    #global samplepath1
   # samplepath1="C:\\Users\\archana.muraly\\Documents\\newfiles"
    #today = str(date.today())
    #print(date.today().strftime('%Y%m%d'))
    #print(datetime.datetime.today().strftime('%Y-%m-%d'))
    # In[3]:
    global keywords
    keywords = pd.read_excel(path+'\\BagOfKeywords_20190403'+'.xlsx',sheet_name='BagOfKeywords') #Update the link
    # #### Extract PDF files
    # In[4]:
   # print(keywords)	
    pdftextFileslist = listtextfiles(samplepath)
    #print(pdftextFileslist)
    pdfoutput_list = []
    pdfkeywordwritten = []
    for samfile in pdftextFileslist:
       # print("entered")
        pdfoutput_list.append(getpdfoutput(samfile))
        #print("OUTPUT-----------------",pdfoutput_list)
        #print("content",getmatchwords(getcontentspdf(samfile)))
        pdfkeywordwritten.append(getmatchwords(getcontentspdf(samfile)))
        #print("Completed:",samfile)
    #list_cleaned = [re.sub(r"[\x83\x86\x85\x84\u2002¾→]", "", file) for file in pdfoutput_list]
    #print("printing outputlist************",len(pdfoutput_list))
   # print("last output",pdfoutput_list)
   # print("key output",pdfkeywordwritten)
    pdftxtfilename = []
    for txt in pdftextFileslist:
        splitpd = os.path.splitext(txt)[0]
        pdftxtfilename.append(os.path.basename(splitpd))
    pdftxtfilepd = pd.DataFrame(pdftxtfilename,columns=(['filename']))
    pdfkeywordpd = pd.DataFrame(pdfkeywordwritten)
    if pdfoutput_list:
        pdfoutputpd = pd.DataFrame(pdfoutput_list)
    else:
        pdfoutputpd=pd.DataFrame([])
    pdftxtfilepdshape = list(pdftxtfilepd.shape)
    #print(len(pdftxtfilepdshape))
    pdfkeywordpdshape = list(pdfkeywordpd.shape)
    #print(len(pdfkeywordpdshape))
    pdfoutputpdshape = list(pdfoutputpd.shape)
    #print(pdfoutputpdshape)
    # #### Organize  extracted text in Dataframe
    # In[90]:
    pdfdf1 = pd.concat([pdftxtfilepd,pdfkeywordpd,pdfoutputpd],axis=1)
    pdfdf1 = pdfdf1.set_index('filename')
    if not pdfdf1.empty:
        pdfa = pdfdf1.iloc[:,0:pdfkeywordpdshape[1]].stack()
        #print(pdfa)
        pdfa_df = pd.DataFrame(pdfa,columns=(['words']))
        pdfb = pdfdf1.iloc[:,pdfoutputpdshape[1]:].stack()
        #print(pdfb)
        pdfb_df = pd.DataFrame(pdfb,columns=(['text']))
        #print(pdfb_df)
        pdfc_df = pd.concat([pdfa_df,pdfb_df],axis=1)
        final_output_pdf = pdfc_df.reset_index()
    else:
        final_output_pdf = pd.DataFrame()
    #final_output.to_csv(samplepath+'\output1.csv')
    # In[92]:
    #print(final_output_pdf)
    # In[93]:
    final_output_pdf.to_csv(samplepath+'\pdf1.csv')
    # ## Word Extraction
    # In[28]:
    wordtextFileslist = listdocfiles(samplepath)
    wordfilenames = []
    wordoutput_list = []
    wordkeywordpd = []
    for samfile in wordtextFileslist:
        print("completed:",samfile)
        splitpd = os.path.splitext(samfile)[0]
        output = (getwordoutput(samfile))
        wordfilenames.append(os.path.basename(splitpd))
        wordoutput_list.append(output)
        wordkeywordpd.append(getwordoutput1(samfile))
    # In[36]:
    wordfilenamespd = pd.DataFrame(wordfilenames,columns=(['filename']))
    wordoutputpd = pd.DataFrame(wordoutput_list)
    wordkeywordpd = pd.DataFrame(wordkeywordpd)
    # In[37]:
    wordfilenamespdshape = list(wordfilenamespd.shape)
    wordoutputpdshape = list(wordoutputpd.shape)
    wordkeywordpdshape = list(wordkeywordpd.shape)
    # #### Organize text in Dataframe
    # In[38]:
    #wordoutputpd
    # In[39]:
    df1 = pd.concat([wordfilenamespd,wordkeywordpd,wordoutputpd],axis=1)
    df1 = df1.set_index('filename')
    if not df1.empty:
        a = df1.iloc[:,0:wordkeywordpdshape[1]].stack()
        a_df = pd.DataFrame(a,columns=(['words']))
        b = df1.iloc[:,wordoutputpdshape[1]:].stack()
        b_df = pd.DataFrame(b,columns=(['text']))
        c_df = pd.concat([a_df,b_df],axis=1)
        word_final_output = c_df.reset_index()
    else:
        word_final_output = pd.DataFrame()
    #final_output.to_csv(r'C:\Madhan\Analytics\Machine_Learning_project_work\AI_ML\Tender_notification\dev\word\fin2.csv')
    # In[40]:
    final_output = pd.concat([final_output_pdf,word_final_output],axis=0)
    # In[94]:
    final_output.head()
    # In[42]:
    final_output.to_csv(samplepath+'\extractions.csv')
    for samfile in wordtextFileslist:
        docxtxt=samfile+'docx.txt'
        
        os.remove(docxtxt)
    return samplepath+'\extractions.csv'
    
    

def listdocfiles(samplepath):
    textFiles = []
    #newpdfFiles = []
    for root, dirs, files in os.walk(samplepath):
        for file in files:
            #print(file)
            if file.endswith(".docx"):
                textFiles.append(os.path.join(root, file))
            elif file.endswith(".DOCX"):
                textFiles.append(os.path.join(root, file))
    return(textFiles)


# In[29]:


def doc2text(filename):
    document = docx.Document(filename)
    para = document
    for line in para.paragraphs:
        with open (os.path.join(samplepath,filename+'docx.txt'),"a",encoding='utf-8') as doc:
            text = doc.write("%s\n" % line.text)
    return('File Written Successfully')


# #### Get Bold contents

# In[30]:


def getbold_contents(filename):
    kwords = list(keywords.words)
    bold_list = []
    document = docx.Document(filename)
    para = document
    #print(kwords)
    for line in document.paragraphs:
        for run in line.runs:
            if run.bold:
                bold_text = run.text.strip()
                if bold_text in kwords:
                    bold_list.append(bold_text)
    for line in para.paragraphs:
        style = line.style.name
        if 'Head' in style:
            header = line.text
            if header in kwords:
                bold_list.append(header)
    content_words = [bd for bd in bold_list if len(bd)>0 and bd.strip()]
    set_content_words = set(content_words)
    if not content_words:
        print('No Extractable information')
    return list(set_content_words)


# In[45]:


def targettxtoutput(filename,content_1):
    sentences = []
    flag=0
    cont_length = len(content_1.split())
    my_file = open(os.path.join(samplepath,filename+'docx.txt'),"r",encoding='utf-8')
    op_filename = samplepath+'\output.txt'
    op_file = open(op_filename,"a",encoding='utf-8')
    c=0
    for line in my_file:
	    if ((content_1 in line) and len(line.split())==cont_length):
		    flag = 1
	    if (len(line)==0 and len(line)<100):
		    flag = 0
	    if (flag==1):
	        sentences.append(line)
    #print(sentences)
    s = []
    for sent in sentences[1:]:
        sentlength = len(sent.split())
        #print(sentlength)
        #print(sent)
        while sentlength == 1 or sentlength == 2:
            sent = sent.strip()
            s.append(sent)
            break
    if s:
        #print(s)
        #print(s[0])
        for sent in sentences:
            if ((content_1 in sent) and len(sent.split())==cont_length):
                c+=1
                flag =2
            if (s[0] in sent):
                flag = 3
            if flag ==2 and c<2:
                #targetsent = sent.strip()
                op_file.write(sent)
    my_file.close()
    op_file.close()
    text = readfromText(op_filename)
    #text = textwrap.fill(text,replace_whitespace=False)
    docxtxt=filename+'docx.txt'
   # os.remove(docxtxt)
    os.remove(op_filename)
    return unidecode(text)


# In[32]:


def getwordoutput(filename):
    doc2text(filename)
    splitpd = os.path.splitext(filename)[0]
    wordfilename = os.path.basename(splitpd)
    output_list = []
    list_cleaned = []
    list_cleaned = []
    keywordwritten = []
    content_words = getbold_contents(filename)
    #print(content_words)
    output_list = []
    for content in content_words:
        output = targettxtoutput(filename,content)
        keywordwritten.append(content)
        output_list.append(output)
        #print(output)
        #with open (os.path.join(samplepath,filename+'_extracted_.txt'),"a",encoding='utf-8') as doc:
         #   text = doc.writelines("%s\n" % output.strip())
    output_list = [out for out in output_list if len(out)>0]
    #list_cleaned = [re.sub(r"[^a-zA-Z.]", " ", file) for file in output_list]
    #list_cleaned1 = [" ".join(file.split()) for file in list_cleaned]
    keywordpd = pd.DataFrame(keywordwritten,columns=(['keywords']))
    #print(len(output_list))
    #outputpd = pd.DataFrame(output_list,columns=(['Extracted_content']))
    #fileoutputpd = pd.concat([keywordpd,outputpd]) 
    return output_list


# In[33]:


def getwordoutput1(filename):
    #doc2text(filename)
    splitpd = os.path.splitext(filename)[0]
    wordfilename = os.path.basename(splitpd)
    output_list = []
    keywordwritten = []
    content_words = getbold_contents(filename)
    #print(content_words)
    output_list = []
    for content in content_words:
        #output = targettxtoutput(filename,content)
        keywordwritten.append(content)
        #output_list.append(output)
        #print(output)
        #with open (os.path.join(samplepath,filename+'_extracted_.txt'),"a",encoding='utf-8') as doc:
         #   text = doc.writelines("%s\n" % output.strip())
    #keywordpd = pd.DataFrame(keywordwritten,columns=(['keywords']))
    #print(len(output_list))
    #outputpd = pd.DataFrame(output_list,columns=(['Extracted_content']))
    #fileoutputpd = pd.concat([keywordpd,outputpd]) 
    return keywordwritten



global stops
stops=['i', 'me', 'my', 'myself', 'we', 
       'our', 'ours', 'ourselves', 'you', 
       "you're", "you've", "you'll", "you'd",
       'your', 'yours', 'yourself', 'yourselves', 
       'he', 'him', 'his', 'himself', 'she', "she's", 
       'her', 'hers', 'herself', 'it', "it's", 'its', 
       'itself', 'they', 'them', 'their', 'theirs', 
       'themselves', 'what', 'which', 'who', 'whom', 
       'this', 'that', "that'll", 'these', 'those', 
       'am', 'is', 'are', 'was', 'were', 'be', 'been', 
       'being', 'have', 'has', 'had', 'having', 'do', 'does',
       'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if',
       'or', 'because', 'as', 'until', 'while', 'of', 'at',
       'by', 'for', 'with', 'about', 'against', 'between', 
       'into', 'through', 'during', 'before', 'after', 
       'above', 'below', 'to', 'from', 'up', 'down',
       'in', 'out', 'on', 'off', 'over', 'under', 'again',
       'further', 'then', 'once', 'here', 'there', 'when',
       'where', 'why', 'how', 'all', 'any', 'both', 'each',
       'few', 'more', 'most', 'other', 'some', 'such', 'no',
       'nor', 'not', 'only', 'own', 'same', 'so', 'than', 
       'too', 'very', 's', 't', 'can', 'will', 'just', 
       'don', "don't", 'should', "should've", 'now', 'd',
       'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't",
       'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't",
       'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 
       'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn', "mustn't",
       'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't",
       'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', 
       "wouldn't"]
global symbol
symbol=['(',')','.']
path1="C:\\Users\\archana.muraly\\Documents\\"
path=r"C:\\Users\\archana.muraly\\Documents\\file2"
extractFile(path)



# In[2]:

